"""
Generate the 3PAO Assessment Methodology document from config data.

Reads config/nist_800_53_controls.json and config/checks/*.json to produce:
  1. docs/assessment-methodology.md
  2. docs/assessment-methodology.docx (if python-docx installed)

Usage:
    cd FedRAMP-SCANNER
    python scripts/generate_methodology_doc.py
"""
from __future__ import annotations

import json
import glob
import os
import textwrap
from pathlib import Path
from datetime import datetime

ROOT = Path(__file__).resolve().parent.parent
CONFIG_DIR = ROOT / "config"
DOCS_DIR = ROOT / "docs"


def _numeric_key(item):
    """Sort key for NIST 800-53 IDs like 'AC', 'AC-2', 'AC-2(1)' — alphabetic family, numeric control."""
    import re
    key = item[0] if isinstance(item, tuple) else item
    m = re.match(r'([A-Z]{2})-?(\d+)?', key)
    if m:
        return (m.group(1), int(m.group(2)) if m.group(2) else 0)
    return (key, 0)
DOCS_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------------------------
# Load data
# ---------------------------------------------------------------------------

def load_controls() -> dict:
    with open(CONFIG_DIR / "nist_800_53_controls.json") as f:
        return json.load(f)

def load_all_checks() -> dict:
    """Return dict keyed by domain code (AC, AU, etc.)."""
    result = {}
    for fpath in sorted(glob.glob(str(CONFIG_DIR / "checks" / "*.json"))):
        with open(fpath) as f:
            data = json.load(f)
        result[data["domain"]] = data
    return result

# ---------------------------------------------------------------------------
# Statistics helpers
# ---------------------------------------------------------------------------

def compute_stats(controls_data, all_checks):
    stats = {
        "total_controls": 0,
        "automated": 0,
        "manual": 0,
        "total_objectives": 0,
        "obj_auto_true": 0,
        "obj_auto_partial": 0,
        "obj_auto_false": 0,
        "total_checks": 0,
        "aws_checks": 0,
        "azure_checks": 0,
        "gcp_checks": 0,
        "doc_requirements": 0,
        "domains": {},
    }

    for fam_id, family in sorted(controls_data["families"].items(), key=_numeric_key):
        domain = family["domain"]
        name = family["name"]
        d = {
            "name": name,
            "controls": 0,
            "automated": 0,
            "manual": 0,
            "objectives": 0,
            "aws": 0,
            "azure": 0,
            "gcp": 0,
        }

        for pid, p in family["controls"].items():
            stats["total_controls"] += 1
            d["controls"] += 1
            is_auto = p.get("automated", False)
            if is_auto:
                stats["automated"] += 1
                d["automated"] += 1
            else:
                stats["manual"] += 1
                d["manual"] += 1

            objs = p.get("objectives", {})
            stats["total_objectives"] += len(objs)
            d["objectives"] += len(objs)
            for obj in objs.values():
                a = obj.get("automatable")
                if a is True:
                    stats["obj_auto_true"] += 1
                elif a == "partial":
                    stats["obj_auto_partial"] += 1
                else:
                    stats["obj_auto_false"] += 1

            if domain in all_checks and pid in all_checks[domain].get("checks", {}):
                pdata = all_checks[domain]["checks"][pid]
                for c in pdata.get("aws", []):
                    stats["aws_checks"] += 1
                    d["aws"] += 1
                for c in pdata.get("azure", []):
                    stats["azure_checks"] += 1
                    d["azure"] += 1
                for c in pdata.get("gcp", []):
                    stats["gcp_checks"] += 1
                    d["gcp"] += 1
                stats["doc_requirements"] += len(
                    pdata.get("objectives_requiring_documentation", [])
                )

        stats["domains"][domain] = d

    stats["total_checks"] = stats["aws_checks"] + stats["azure_checks"] + stats["gcp_checks"]
    return stats


# ---------------------------------------------------------------------------
# Markdown generation
# ---------------------------------------------------------------------------

def generate_markdown(controls_data, all_checks, stats) -> str:
    lines = []
    w = lines.append

    # --- Title & Frontmatter ---
    w("# FedRAMP Cloud Compliance Scanner — Assessment Methodology")
    w("")
    w("**Document Classification:** For Official Use — Assessment Staff Only")
    w("")
    w(f"**Version:** 1.0 | **Date:** {datetime.now().strftime('%B %d, %Y')} | **Author:** Securitybricks (3PAO, powered by Aprio)")
    w("")
    w("---")
    w("")

    # --- Table of Contents ---
    w("## Table of Contents")
    w("")
    w("1. [Executive Summary](#1-executive-summary)")
    w("2. [Purpose and Audience](#2-purpose-and-audience)")
    w("3. [Authoritative Sources and Traceability](#3-authoritative-sources-and-traceability)")
    w("4. [Assessment Architecture](#4-assessment-architecture)")
    w("   - 4.1 [Check-to-Objective Mapping](#41-check-to-objective-mapping)")
    w("   - 4.2 [Three-Tier Evaluation Model](#42-three-tier-evaluation-model)")
    w("   - 4.3 [Cloud Provider API Baselines](#43-cloud-provider-api-baselines)")
    w(f"   - 4.4 [Why Only {stats['automated']} of {stats['total_controls']} Controls Are Automated](#44-why-only-{stats['automated']}-of-{stats['total_controls']}-controls-are-automated)")
    w("5. [Coverage Matrix Summary](#5-coverage-matrix-summary)")
    w("   - 5.1 [Overall Statistics](#51-overall-statistics)")
    w("   - 5.2 [Domain-Level Coverage](#52-domain-level-coverage)")
    w("   - 5.3 [Objective Automatable Classification](#53-objective-automatable-classification)")
    w("6. [Complete Control Reference](#6-complete-control-reference)")

    # TOC for each domain
    for fam_idx, (fam_id, family) in enumerate(sorted(controls_data["families"].items(), key=_numeric_key)):
        domain = family["domain"]
        name = family["name"]
        w(f"   - 6.{fam_idx+1} [{domain} — {name}](#{domain.lower()}--{name.lower().replace(' ', '-').replace('&', '').replace('  ', '-')})")

    w("7. [3PAO Manual Assessment Guide](#7-cca-manual-assessment-guide)")
    w("   - 7.1 [How to Use This Guide](#71-how-to-use-this-guide)")
    w("   - 7.2 [Manual Control Reference](#72-manual-control-reference)")
    w("8. [Appendix A — API Call Reference](#8-appendix-a--api-call-reference)")
    w("9. [Appendix B — Glossary](#9-appendix-b--glossary)")
    w("")
    w("---")
    w("")

    # --- Section 1: Executive Summary ---
    w("## 1. Executive Summary")
    w("")
    w("The FedRAMP Cloud Compliance Scanner is an automated assessment tool built by Securitybricks, a FedRAMP Third-Party Assessment Organization (3PAO) powered by Aprio. It evaluates Defense Industrial Base (DIB) contractor cloud environments against FedRAMP requirements by querying cloud service provider (CSP) APIs and comparing configuration states to NIST SP 800-53 Rev 5 security controls.")
    w("")
    w("This document serves as the **authoritative methodology reference** for FedRAMP Assessors (3PAOs) using the scanner. It explains:")
    w("")
    w("- **How** each of the 110 NIST 800-53 Rev 5 controls is evaluated")
    w("- **Which** cloud APIs are queried and what constitutes a passing or failing check")
    w("- **Why** each check maps to specific NIST SP 800-53A assessment objectives")
    w("- **What** 3PAOs must do for the 39 controls that require manual assessment")
    w("- **Where** the authoritative sources and traceability chain originates")
    w("")
    w(f"The scanner implements **{stats['total_checks']} cloud-specific technical checks** across AWS ({stats['aws_checks']}), Azure ({stats['azure_checks']}), and GCP ({stats['gcp_checks']}), mapped to **{stats['total_objectives']} NIST SP 800-53A assessment objectives** across all 110 controls and 14 FedRAMP control families.")
    w("")
    w("---")
    w("")

    # --- Section 2: Purpose and Audience ---
    w("## 2. Purpose and Audience")
    w("")
    w("### Who Should Read This Document")
    w("")
    w("| Role | How to Use This Document |")
    w("|------|------------------------|")
    w("| **Lead 3PAO** | Validate the scanner's methodology against 800-53A before accepting automated results |")
    w("| **3PAO (Technical)** | Reference during assessment to understand what each check evaluates and which API responses constitute evidence |")
    w("| **3PAO (Policy/Process)** | Use Section 7 as a structured guide for manual control assessments — interview questions, evidence artifacts, and determination criteria |")
    w("| **Assessment Team Lead** | Review coverage matrix to understand which objectives are automated vs. manual |")
    w("| **Quality Assurance** | Verify traceability from check results back to 800-53A objectives |")
    w("")
    w("### How This Document Builds Trust")
    w("")
    w("For a 3PAO to rely on automated tool results in a FedRAMP assessment, they need to verify:")
    w("")
    w("1. **Traceability** — Every automated check traces back to a specific NIST SP 800-53A assessment objective")
    w("2. **Completeness** — The tool identifies which objectives it covers and which require manual assessment")
    w("3. **Accuracy** — Checks query the correct cloud APIs and evaluate the right configuration properties")
    w("4. **Transparency** — The methodology is fully documented, not a black box")
    w("")
    w("This document satisfies all four requirements.")
    w("")
    w("---")
    w("")

    # --- Section 3: Authoritative Sources ---
    w("## 3. Authoritative Sources and Traceability")
    w("")
    w("The scanner's check library is derived from and traceable to the following authoritative sources:")
    w("")
    w("| Source | Version | Purpose | Reference |")
    w("|--------|---------|---------|-----------|")
    w("| **NIST SP 800-53 Rev 5** | Feb 2020 | 110 security controls across 14 families | [csrc.nist.gov](https://csrc.nist.gov/publications/detail/sp/800-53 Rev 5/rev-2/final) |")
    w("| **NIST SP 800-53A** | Jun 2018 | 319 assessment objectives (\"determine if\" statements) | [csrc.nist.gov](https://csrc.nist.gov/publications/detail/sp/800-53 Rev 5a/final) |")
    w("| **NIST SP 800-172** | Feb 2021 | Enhanced security controls for Level 3 | [csrc.nist.gov](https://csrc.nist.gov/publications/detail/sp/800-172/final) |")
    w("| **FAR 52.204-21** | 2016 | 17 basic safeguarding controls for Level 1 | [acquisition.gov](https://www.acquisition.gov/far/52.204-21) |")
    w("| **FedRAMP Model** | Dec 2021 | Three-level certification model | [dodcio.defense.gov](https://dodcio.defense.gov/FedRAMP/) |")
    w("| **AWS Config Rules** | Current | ~200 rules mapped to NIST 800-53 Rev 5 | [docs.aws.amazon.com](https://docs.aws.amazon.com/config/latest/developerguide/operational-best-controls-for-nist_800-53 Rev 5.html) |")
    w("| **Azure Policy** | Current | ~200 policy definitions for NIST 800-53 Rev 5 R2 | [learn.microsoft.com](https://learn.microsoft.com/en-us/azure/governance/policy/samples/nist-sp-800-53 Rev 5-r2) |")
    w("| **GCP CIS Benchmark** | Current | GCP security controls aligned to NIST controls | [cloud.google.com](https://cloud.google.com/security/compliance/cis-benchmarks) |")
    w("")
    w("### Traceability Chain")
    w("")
    w("Every check in the scanner traces back through the following chain:")
    w("")
    w("| Step | Stage | Description |")
    w("|------|-------|-------------|")
    w("| 1 | **FedRAMP Level** | L1 / L2 / L3 certification tier |")
    w("| 2 | **NIST SP 800-53 Rev 5 Control** | One of 110 security requirements |")
    w("| 3 | **800-53A Assessment Objective** | Specific \"determine if\" statement |")
    w("| 4 | **Scanner Check** | Cloud-specific configuration test |")
    w("| 5 | **Cloud API Call** | Read-only query to AWS, Azure, or GCP |")
    w("| 6 | **Compliance Determination** | Met / Not Met / Manual Review |")
    w("")
    w("Every finding in the scanner report traces back through this chain to the authoritative NIST standard.")
    w("")
    w("---")
    w("")

    # --- Section 4: Assessment Architecture ---
    w("## 4. Assessment Architecture")
    w("")
    w("### 4.1 Check-to-Objective Mapping")
    w("")
    w("NIST SP 800-53A defines **319 assessment objectives** across the 110 NIST SP 800-53 Rev 5 controls. Each objective is a discrete \"determine if\" statement that an assessor must evaluate.")
    w("")
    w("The scanner maps every automated check to the specific 800-53A objectives it supports via the `supports_objectives` field. For example:")
    w("")
    w("```json")
    w("{")
    w('  "check_id": "ac-3.1.1-aws-001",')
    w('  "name": "Root account access keys disabled",')
    w('  "supports_objectives": ["[d]"],')
    w('  "service": "IAM",')
    w('  "api_call": "iam.get_account_summary",')
    w('  "severity": "critical"')
    w("}")
    w("")
    w("```")
    w("")
    w("**Objective [d]** for control 3.1.1 states: *\"system access is limited to authorized users.\"* Disabling root account access keys directly enforces this by preventing the most privileged account from using long-term credentials.")
    w("")
    w("This mapping enables:")
    w("- **Per-objective coverage scoring** — the report shows which objectives are covered by automated checks, which require documentation, and which are not tested")
    w("- **Gap identification** — 3PAOs can immediately see which objectives need manual verification")
    w("- **Audit traceability** — every Met/Not Met determination links to specific 800-53A language")
    w("")
    w("### 4.2 Three-Tier Evaluation Model")
    w("")
    w("The scanner classifies every 800-53A assessment objective into one of three tiers:")
    w("")
    w("| Tier | Classification | Count | Description |")
    w("|------|---------------|-------|-------------|")
    w(f"| **Tier 1** | Fully Automatable | {stats['obj_auto_true']} | Cloud API configuration check provides a definitive Met/Not Met determination |")
    w(f"| **Tier 2** | Partially Automatable | {stats['obj_auto_partial']} | Cloud API provides supporting evidence, but 3PAO must verify organizational context |")
    w(f"| **Tier 3** | Not Automatable | {stats['obj_auto_false']} | Requires documentation review, interviews, or physical inspection |")
    w("")
    w("**Tier 1 — Fully Automatable:** The API response alone determines compliance. Example: *\"MFA is enabled for all console users\"* — the credential report provides a binary yes/no.")
    w("")
    w("**Tier 2 — Partially Automatable:** The API response provides evidence that supports a determination, but the 3PAO must also verify organizational context. Example: *\"Authorized users are identified\"* — IAM user lists show WHO has access, but the 3PAO must verify this matches the organization's authorized user roster.")
    w("")
    w("**Tier 3 — Not Automatable:** No cloud API can evaluate this objective. Example: *\"Visitors are escorted\"* (physical security) or *\"Security awareness training is provided\"* (organizational process).")
    w("")
    w("### 4.3 Cloud Provider API Baselines")
    w("")
    w("The scanner uses read-only API calls across three cloud service providers. All access is via read-only IAM roles — **no configuration changes are ever made** to the client's environment.")
    w("")
    w("| Provider | Access Method | Permissions Required | Checks |")
    w("|----------|--------------|---------------------|--------|")
    w(f"| **AWS** | STS AssumeRole (cross-account) | `SecurityAudit` + `ViewOnlyAccess` managed policies | {stats['aws_checks']} |")
    w(f"| **Azure** | Service Principal (ClientSecretCredential) | `Reader` + `Security Reader` roles + Microsoft Graph API | {stats['azure_checks']} |")
    w(f"| **GCP** | Service Account (JSON key) | `Viewer` + `Security Reviewer` + `Security Center Admin` roles | {stats['gcp_checks']} |")
    w("")

    # List key APIs per provider
    w("**Key AWS Services Queried:** IAM, STS, EC2, VPC, S3, CloudTrail, CloudWatch, Config, GuardDuty, Security Hub, KMS, SSM, Inspector, WAFv2, ELB, CloudFront, RDS, Organizations, ACM, Route 53, Network Firewall, DynamoDB, API Gateway, CodePipeline, Athena, SNS, ECR, Health")
    w("")
    w("**Key Azure Services Queried:** Entra ID (Graph API), Network, Compute, Storage, Key Vault, Security Center, Monitor, Policy, Authorization, SQL, App Service, Sentinel, Advisor, Automation, Resource Graph, Guest Configuration")
    w("")
    w("**Key GCP Services Queried:** IAM, Cloud Resource Manager, Compute, VPC, Storage, Cloud KMS, Cloud Logging, Cloud Monitoring, Security Command Center, OS Config, Binary Authorization, Container Analysis, Web Security Scanner, Cloud SQL, BigQuery, Cloud DNS, Recommender, Cloud IDS, Cloud Armor, BeyondCorp, Organization Policy")
    w("")
    # --- Section 4.4: Why Only N Controls Are Automated ---
    w(f"### 4.4 Why Only {stats['automated']} of {stats['total_controls']} Controls Are Automated")
    w("")
    w(f"NIST SP 800-53 Rev 5 defines {stats['total_controls']} security controls, but only {stats['automated']} "
      f"({stats['automated']*100//stats['total_controls']}%) can be meaningfully evaluated through cloud API queries. "
      f"This is not a gap in scanner coverage — it reflects the fundamental nature of the controls themselves. "
      f"The remaining {stats['manual']} controls govern activities that occur outside of cloud infrastructure "
      f"and cannot be observed through any API.")
    w("")
    w("**Controls That Cannot Be Automated:**")
    w("")
    w("| Category | Families | Controls | Why Not Automatable |")
    w("|----------|----------|----------|---------------------|")

    # Get per-domain stats
    domain_order = ["AC", "AT", "AU", "CA", "CM", "CP", "IA", "IR", "MA", "MP",
                    "PE", "PL", "PM", "PS", "PT", "RA", "SA", "SC", "SI", "SR"]
    domain_stats = {}
    for fam_id, family in controls_data["families"].items():
        d = family["domain"]
        if d not in domain_stats:
            domain_stats[d] = {"controls": 0, "automated": 0, "manual": 0}
        for pid, p in family["controls"].items():
            domain_stats[d]["controls"] += 1
            if p.get("automated", False):
                domain_stats[d]["automated"] += 1
            else:
                domain_stats[d]["manual"] += 1

    pe_c = domain_stats.get("PE", {}).get("controls", 0)
    pm_c = domain_stats.get("PM", {}).get("controls", 0)
    ps_c = domain_stats.get("PS", {}).get("controls", 0)
    at_c = domain_stats.get("AT", {}).get("controls", 0)
    zero_total = pe_c + pm_c + ps_c + at_c

    w(f"| Physical & Environmental | PE ({pe_c} controls) | {pe_c} | Facility access, environmental protections, fire suppression — no API can verify a locked door |")
    w(f"| Organizational Governance | PM ({pm_c} controls) | {pm_c} | Risk strategy, authorization processes, insider threat programs — organizational policies, not cloud configs |")
    w(f"| Personnel Security | PS ({ps_c} controls) | {ps_c} | Background checks, screening, access agreements, termination — HR processes requiring human verification |")
    w(f"| Awareness & Training | AT ({at_c} controls) | {at_c} | Security training programs, role-based training — requires review of materials and completion records |")
    w(f"| **Subtotal** | **4 families (0% automation)** | **{zero_total}** | **{zero_total*100//stats['total_controls']}% of the framework** |")
    w("")
    w("**Policy & Procedure Objectives Within Automated Families:**")
    w("")
    w("Even in families with automated checks, many controls ask: *\"Does the organization define and document a policy/procedure for X?\"* "
      "These Tier 3 objectives require a 3PAO to review written policies and standard operating procedures — artifacts "
      "that exist as documents, not as cloud API states. This is why families like Access Control (AC) show only "
      f"{domain_stats.get('AC', {}).get('automated', 0)}/{domain_stats.get('AC', {}).get('controls', 0)} automation: "
      "the automated controls check *configuration states* (MFA, least-privilege roles, session timeouts), "
      "while the manual controls verify that policies, procedures, and approval workflows exist and are followed.")
    w("")
    w("**What *Can* Be Automated:**")
    w("")
    w(f"The {stats['automated']} automated controls share a common trait: their compliance state is observable "
      "through cloud provider APIs as a **configuration property** that is either present or absent. Examples:")
    w("")
    w("- **AC-6(3):** *Is privileged access restricted to specific accounts?* → Check IAM policies for least-privilege roles (API-verifiable)")
    w("- **AU-6:** *Are audit records reviewed and analyzed?* → Check CloudTrail/Activity Log enabled and forwarded (API-verifiable)")
    w("- **SC-8:** *Is transmitted information protected?* → Check TLS/SSL configuration on load balancers (API-verifiable)")
    w("- **PE-3:** *Is physical access controlled?* → Requires on-site inspection of badge readers and guard stations (not API-verifiable)")
    w("")
    w(f"> **Bottom line:** The {stats['automated']}/{stats['total_controls']} automation ratio is inherent to the "
      "NIST 800-53 framework, which intentionally covers physical, procedural, and organizational security alongside "
      "technical controls. A scanner claiming 100% automation of 800-53 would be misrepresenting what the framework requires.")
    w("")
    w("---")
    w("")

    # --- Section 5: Coverage Matrix ---
    w("## 5. Coverage Matrix Summary")
    w("")
    w("### 5.1 Overall Statistics")
    w("")
    w("| Metric | Value |")
    w("|--------|-------|")
    w(f"| NIST 800-53 Rev 5 Controls | {stats['total_controls']} |")
    w(f"| NIST 800-53A Assessment Objectives | {stats['total_objectives']} |")
    w(f"| Controls with Automated Checks | {stats['automated']} ({stats['automated']*100//stats['total_controls']}%) |")
    w(f"| Controls Requiring Manual Assessment | {stats['manual']} ({stats['manual']*100//stats['total_controls']}%) |")
    w(f"| Total Cloud-Specific Technical Checks | {stats['total_checks']} |")
    w(f"| AWS Checks | {stats['aws_checks']} |")
    w(f"| Azure Checks | {stats['azure_checks']} |")
    w(f"| GCP Checks | {stats['gcp_checks']} |")
    w(f"| Documentation Evidence Requirements | {stats['doc_requirements']} |")
    w("")

    w("### 5.2 Domain-Level Coverage")
    w("")
    w("The table below shows the scanner's coverage across all 14 FedRAMP control families. Each domain is broken down by the number of NIST 800-53 Rev 5 controls, how many are automated vs. manual, the total 800-53A assessment objectives, and the cloud-specific checks implemented for each provider. The **Automation Rate** shows the percentage of controls in each domain that are fully automated by the scanner.")
    w("")
    w("| Domain | Name | Controls | Automated | Manual | Objectives | AWS | Azure | GCP | Automation Rate |")
    w("|--------|------|-----------|-----------|--------|------------|-----|-------|-----|-----------------|")

    domain_order = ["AC", "AT", "AU", "CM", "IA", "IR", "MA", "MP", "PE", "PS", "RA", "CA", "SC", "SI"]
    for domain_code in domain_order:
        if domain_code in stats["domains"]:
            d = stats["domains"][domain_code]
            auto_pct = round(d["automated"] * 100 / d["controls"]) if d["controls"] > 0 else 0
            w(f"| {domain_code} | {d['name']} | {d['controls']} | {d['automated']} | {d['manual']} | {d['objectives']} | {d['aws']} | {d['azure']} | {d['gcp']} | {auto_pct}% |")

    total_auto_pct = round(stats['automated'] * 100 / stats['total_controls']) if stats['total_controls'] > 0 else 0
    w(f"| **Total** | | **{stats['total_controls']}** | **{stats['automated']}** | **{stats['manual']}** | **{stats['total_objectives']}** | **{stats['aws_checks']}** | **{stats['azure_checks']}** | **{stats['gcp_checks']}** | **{total_auto_pct}%** |")
    w("")

    w("### 5.3 Objective Automatable Classification")
    w("")
    w("| Classification | Count | Percentage | Scanner Handling |")
    w("|---------------|-------|------------|-----------------|")
    w(f"| Fully Automatable | {stats['obj_auto_true']} | {stats['obj_auto_true']*100//stats['total_objectives']}% | Automated check provides Met/Not Met determination |")
    w(f"| Partially Automatable | {stats['obj_auto_partial']} | {stats['obj_auto_partial']*100//stats['total_objectives']}% | Automated check provides evidence; 3PAO verifies context |")
    w(f"| Not Automatable | {stats['obj_auto_false']} | {stats['obj_auto_false']*100//stats['total_objectives']}% | Flagged as Documentation Required; 3PAO assesses manually |")
    w("")
    w("---")
    w("")

    # --- Section 6: Complete Control Reference (was Section 7) ---
    w("## 6. Complete Control Reference")
    w("")
    w("This section provides the complete technical reference for every NIST SP 800-53 Rev 5 control, organized by FedRAMP family. For each control, it shows:")
    w("")
    w("- The requirement text and FedRAMP baseline")
    w("- All NIST SP 800-53A assessment objectives")
    w("- Cloud-specific automated checks with API calls, services, and severity")
    w("- Objective mapping (which checks support which objectives)")
    w("- Documentation requirements for non-automatable objectives")
    w("")

    for fam_id, family in sorted(controls_data["families"].items(), key=_numeric_key):
        domain = family["domain"]
        name = family["name"]
        d = stats["domains"].get(domain, {})

        w(f"### {domain} — {name}")
        w("")
        w(f"**Controls:** {d.get('controls', 0)} | **Automated:** {d.get('automated', 0)} | **Manual:** {d.get('manual', 0)} | **Objectives:** {d.get('objectives', 0)} | **Checks:** AWS {d.get('aws', 0)}, Azure {d.get('azure', 0)}, GCP {d.get('gcp', 0)}")
        w("")

        for pid, p in sorted(family["controls"].items(), key=_numeric_key):
            level = p.get("level", "L2")
            is_auto = p.get("automated", False)
            auto_label = "Automated" if is_auto else "Manual"
            objs = p.get("objectives", {})

            w(f"#### {pid} — {p['requirement']}")
            w("")
            w(f"**Level:** {level} | **Type:** {auto_label} | **Objectives:** {len(objs)}")
            if p.get("far_mapping"):
                w(f" | **FAR 52.204-21:** {p['far_mapping']}")
            w("")

            # Assessment objectives table
            if objs:
                w("**Assessment Objectives:**")
                w("")
                w("| ID | Objective | Automatable |")
                w("|----|-----------|-------------|")
                for oid, obj in sorted(objs.items()):
                    a = obj.get("automatable")
                    if a is True:
                        a_label = "Yes"
                    elif a == "partial":
                        a_label = "Partial"
                    else:
                        a_label = "No"
                    w(f"| {pid}{oid} | {obj['text'][:120]} | {a_label} |")
                w("")

            # Automated checks
            if domain in all_checks and pid in all_checks[domain].get("checks", {}):
                pdata = all_checks[domain]["checks"][pid]

                has_checks = False
                for cloud in ["aws", "azure", "gcp"]:
                    if pdata.get(cloud):
                        has_checks = True
                        break

                if has_checks:
                    w("**Automated Checks:**")
                    w("")
                    w("| Check ID | Cloud | Name | Service | API Call | Severity | Objectives |")
                    w("|----------|-------|------|---------|---------|----------|------------|")
                    for cloud in ["aws", "azure", "gcp"]:
                        for c in pdata.get(cloud, []):
                            so = ", ".join(c.get("supports_objectives", []))
                            api = c.get("api_call", "N/A")
                            w(f"| `{c['check_id']}` | {cloud.upper()} | {c['name']} | {c.get('service', '')} | `{api}` | {c.get('severity', '')} | {so} |")
                    w("")

                # Documentation requirements
                doc_reqs = pdata.get("objectives_requiring_documentation", [])
                if doc_reqs:
                    w("**Documentation Requirements:**")
                    w("")
                    for d_req in doc_reqs:
                        w(f"- **{pid}{d_req['id']}**: {d_req['text']} — *{d_req.get('evidence_needed', 'Evidence required')}*")
                    w("")

            # Manual guidance
            if not is_auto and p.get("manual_guidance"):
                w(f"**3PAO Manual Assessment Guidance:** {p['manual_guidance']}")
                w("")

            w("")

    w("---")
    w("")

    # --- Section 7: 3PAO Manual Assessment Guide (was Section 6) ---
    w("## 7. 3PAO Manual Assessment Guide")
    w("")
    w("### 7.1 How to Use This Guide")
    w("")
    w("For the 39 controls classified as **Manual Review Required**, the scanner cannot make an automated determination. The 3PAO must independently assess these controls using the guidance below.")
    w("")
    w("For each manual control, this guide provides:")
    w("")
    w("1. **Assessment Objectives** — The exact 800-53A \"determine if\" statements the 3PAO must evaluate")
    w("2. **Assessment Guidance** — Specific steps, interview topics, and configuration areas to examine")
    w("3. **Evidence Artifacts** — Documents, records, and artifacts the 3PAO should request from the OSC")
    w("4. **Determination Criteria** — What constitutes a Met vs. Not Met finding")
    w("")
    w("**Note:** Some \"manual\" controls have automated checks that provide *supporting evidence* (e.g., cloud configurations). These checks do not determine compliance but give the 3PAO baseline data to inform their manual assessment.")
    w("")

    w("### 7.2 Manual Control Reference")
    w("")

    # Build manual controls
    for fam_id, family in sorted(controls_data["families"].items(), key=_numeric_key):
        domain = family["domain"]
        name = family["name"]

        manual_controls = {
            pid: p
            for pid, p in sorted(family["controls"].items(), key=_numeric_key)
            if not p.get("automated", False)
        }
        if not manual_controls:
            continue

        w(f"#### {domain} — {name} (Manual Controls)")
        w("")

        for pid, p in sorted(manual_controls.items(), key=_numeric_key):
            w(f"##### Control {pid}: {p['requirement']}")
            w("")
            w(f"**FedRAMP Baseline:** {p.get('level', 'L2')} | **Domain:** {domain}")
            w("")

            # Assessment Objectives
            objs = p.get("objectives", {})
            if objs:
                w("**Assessment Objectives (NIST SP 800-53A):**")
                w("")
                for oid, obj in sorted(objs.items()):
                    w(f"- **{pid}{oid}**: Determine if {obj['text']}")
                w("")

            # Assessment Guidance
            mg = p.get("manual_guidance", "")
            if mg:
                w("**3PAO Assessment Guidance:**")
                w("")
                w(f"> {mg}")
                w("")

            # Evidence needed from doc requirements
            if domain in all_checks and pid in all_checks[domain].get("checks", {}):
                pdata = all_checks[domain]["checks"][pid]
                doc_reqs = pdata.get("objectives_requiring_documentation", [])
                if doc_reqs:
                    w("**Evidence Artifacts to Request:**")
                    w("")
                    w("| Objective | Evidence Needed |")
                    w("|-----------|----------------|")
                    for d in doc_reqs:
                        w(f"| {pid}{d['id']} | {d.get('evidence_needed', 'Documentation required')} |")
                    w("")

                # Supporting automated checks
                has_supporting = False
                for cloud in ["aws", "azure", "gcp"]:
                    if pdata.get(cloud):
                        has_supporting = True
                        break
                if has_supporting:
                    w("**Supporting Automated Checks** (provide baseline data for 3PAO review):")
                    w("")
                    w("| Cloud | Check | API Call | What It Evaluates |")
                    w("|-------|-------|---------|-------------------|")
                    for cloud in ["aws", "azure", "gcp"]:
                        for c in pdata.get(cloud, []):
                            w(f"| {cloud.upper()} | {c['check_id']} | `{c.get('api_call', 'N/A')}` | {c['name']} |")
                    w("")

            # Determination criteria
            w("**Determination Criteria:**")
            w("")
            w(f"- **Met:** All {len(objs)} assessment objectives are satisfied with documented evidence")
            w(f"- **Not Met:** One or more objectives lack sufficient evidence or are demonstrably not implemented")
            w("")
            w("---")
            w("")

    w("")
    w("---")
    w("")

    # --- Appendix A: API Call Reference ---
    w("## 8. Appendix A — API Call Reference")
    w("")
    w("Complete list of unique cloud API calls made by the scanner, organized by provider.")
    w("")

    for cloud_label, cloud_key in [("AWS", "aws"), ("Azure", "azure"), ("GCP", "gcp")]:
        api_calls = set()
        services = set()
        for domain_code, domain_data in all_checks.items():
            for pid, pdata in domain_data.get("checks", {}).items():
                for c in pdata.get(cloud_key, []):
                    api = c.get("api_call", "")
                    svc = c.get("service", "")
                    if api:
                        api_calls.add(api)
                    if svc:
                        services.add(svc)

        w(f"### {cloud_label} ({len(api_calls)} unique API calls across {len(services)} services)")
        w("")
        w("| Service | API Call | Used By (Controls) |")
        w("|---------|---------|---------------------|")

        # Build api -> controls mapping
        api_controls = {}
        for domain_code, domain_data in all_checks.items():
            for pid, pdata in domain_data.get("checks", {}).items():
                for c in pdata.get(cloud_key, []):
                    api = c.get("api_call", "")
                    svc = c.get("service", "")
                    if api:
                        key = (svc, api)
                        if key not in api_controls:
                            api_controls[key] = set()
                        api_controls[key].add(pid)

        for (svc, api), pids in sorted(api_controls.items()):
            pids_str = ", ".join(sorted(pids, key=_numeric_key)[:5])
            if len(pids) > 5:
                pids_str += f" (+{len(pids)-5} more)"
            w(f"| {svc} | `{api}` | {pids_str} |")

        w("")

    w("---")
    w("")

    # --- Appendix B: Glossary ---
    w("## 9. Appendix B — Glossary")
    w("")
    w("| Term | Definition |")
    w("|------|-----------|")
    w("| **3PAO** | Certified FedRAMP Assessor — individual authorized to conduct FedRAMP assessments |")
    w("| **3PAO** | FedRAMP Third-Party Assessment Organization — accredited organization that employs 3PAOs |")
    w("| **OSC** | Organization Seeking Certification — the CSP being assessed |")
    w("| **CUI** | Controlled Unclassified Information — sensitive government information requiring protection |")
    w("| **FCI** | Federal Contract Information — information provided by or generated for the government under contract |")
    w("| **DIB** | Defense Industrial Base — companies that supply products/services to the DoD |")
    w("| **CSP** | Cloud Service Provider — AWS, Azure, or GCP |")
    w("| **Met** | The control/objective is fully implemented based on automated or manual evidence |")
    w("| **Not Met** | The control/objective is not implemented or has deficiencies |")
    w("| **Manual Review** | The control requires 3PAO manual assessment — cannot be determined by automated checks alone |")
    w("| **Assessment Objective** | A specific \"determine if\" statement from NIST SP 800-53A that must be evaluated |")
    w("| **POA&M** | Plan of Action and Milestones — remediation plan for Not Met findings |")
    w("| **SSP** | System Security Plan — document describing the system boundary, environment, and security controls |")
    w("| **STS** | Security Token Service — AWS service for assuming cross-account roles |")
    w("| **IAM** | Identity and Access Management — cloud service for managing users, roles, and permissions |")
    w("")
    w("---")
    w("")
    w("## Document Information")
    w("")
    w("This methodology reference is auto-generated from the scanner's configuration files (`config/nist_800_53_controls.json` and `config/checks/*.json`). All check definitions, objective mappings, and coverage data are derived directly from the scanner's authoritative data sources.")
    w("")
    w("For the interactive version of this document, see the **Assessment Methodology** tab in the scanner's Help blade.")
    w("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def generate_docx(md_path: Path, docx_path: Path):
    """Convert the markdown to a .docx file using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        print("Skipping .docx generation.")
        return

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # Navy

    # Read markdown and convert to docx
    with open(md_path) as f:
        lines = f.readlines()

    current_table_rows = []
    in_table = False
    in_code_block = False
    code_lines = []

    def flush_table():
        nonlocal current_table_rows, in_table
        if not current_table_rows:
            return
        # Parse table
        headers = [h.strip() for h in current_table_rows[0].strip("|").split("|")]
        data_rows = []
        for row in current_table_rows[2:]:  # Skip header and separator
            cells = [c.strip() for c in row.strip("|").split("|")]
            data_rows.append(cells)

        cols = len(headers)
        table = doc.add_table(rows=1 + len(data_rows), cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h.replace("**", "")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for r_idx, row_data in enumerate(data_rows):
            for c_idx in range(min(len(row_data), cols)):
                cell = table.rows[r_idx + 1].cells[c_idx]
                text = row_data[c_idx].replace("**", "").replace("`", "")
                cell.text = text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()  # spacing
        current_table_rows = []
        in_table = False

    def flush_code():
        nonlocal code_lines, in_code_block
        if code_lines:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            code_lines = []
        in_code_block = False

    for line in lines:
        stripped = line.rstrip("\n")

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(stripped)
            continue

        # Table rows
        if stripped.startswith("|") and "|" in stripped[1:]:
            if not in_table:
                in_table = True
                current_table_rows = []
            current_table_rows.append(stripped)
            continue
        else:
            if in_table:
                flush_table()

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:].strip(), level=0)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            continue
        if stripped.startswith("##### "):
            doc.add_heading(stripped[6:].strip(), level=4)
            continue

        # Horizontal rules
        if stripped == "---":
            continue

        # Empty lines
        if not stripped:
            continue

        # Bullet points
        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, text)
            continue

        # Blockquotes
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(stripped[2:])
            run.italic = True
            continue

        # Regular paragraphs
        p = doc.add_paragraph()
        _add_formatted_run(p, stripped)

    # Flush remaining
    if in_table:
        flush_table()
    if in_code_block:
        flush_code()

    doc.save(str(docx_path))
    print(f"Generated: {docx_path}")


def _add_formatted_run(paragraph, text):
    """Add text to a paragraph with basic bold/code formatting."""
    from docx.shared import Pt, RGBColor
    import re

    # Simple parsing of **bold** and `code`
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("Loading configuration data...")
    controls_data = load_controls()
    all_checks = load_all_checks()

    print("Computing statistics...")
    stats = compute_stats(controls_data, all_checks)

    print(f"  Controls: {stats['total_controls']}")
    print(f"  Objectives: {stats['total_objectives']}")
    print(f"  Checks: {stats['total_checks']} (AWS={stats['aws_checks']}, Azure={stats['azure_checks']}, GCP={stats['gcp_checks']})")
    print(f"  Doc requirements: {stats['doc_requirements']}")

    print("\nGenerating Markdown document...")
    md_content = generate_markdown(controls_data, all_checks, stats)
    md_path = DOCS_DIR / "assessment-methodology.md"
    with open(md_path, "w") as f:
        f.write(md_content)
    print(f"Generated: {md_path} ({len(md_content):,} characters)")

    print("\nGenerating DOCX document...")
    docx_path = DOCS_DIR / "assessment-methodology.docx"
    generate_docx(md_path, docx_path)

    print("\nDone!")


if __name__ == "__main__":
    main()
